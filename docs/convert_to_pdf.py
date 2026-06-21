"""Convert paper.md to PDF via DOCX using python-docx."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x1A, 0x3C, 0x6E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
BASE = Path(__file__).parent


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE


def add_body(doc, text, bold=False, size=11, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def main():
    md = (BASE / "paper.md").read_text(encoding="utf-8")
    
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    sections = re.split(r'\n## ', md)
    
    first = sections[0]
    lines = first.strip().split("\n")
    for line in lines:
        line = line.strip()
        if line.startswith("# "):
            title = doc.add_heading(line[2:].strip(), 0)
            for run in title.runs:
                run.font.color.rgb = BLUE
        elif line.startswith("**Author"):
            add_body(doc, line.replace("**", ""), bold=True, size=12)
        elif line.startswith("**"):
            clean = line.replace("**", "")
            add_body(doc, clean, bold=True)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(line.lstrip("> ").strip())
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = DARK
        elif line == "---":
            continue
        elif line == "":
            continue
        else:
            clean = line.replace("**", "")
            add_body(doc, clean)

    for section in sections[1:]:
        lines = section.strip().split("\n")
        heading_text = lines[0].strip()
        heading_clean = re.sub(r'^\d+\.\s*', '', heading_text)
        add_heading(doc, heading_clean, 1)

        i = 1
        while i < len(lines):
            line = lines[i].strip()

            if line.startswith("### "):
                add_heading(doc, line.lstrip("# ").strip(), 2)
                i += 1
                continue

            if line.startswith("#### "):
                add_heading(doc, line.lstrip("# ").strip(), 3)
                i += 1
                continue

            if line.startswith("> "):
                quote_lines = []
                while i < len(lines) and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                    text = lines[i].strip().lstrip("> ").strip()
                    if text:
                        quote_lines.append(text)
                    elif quote_lines:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1.5)
                        run = p.add_run("\n".join(quote_lines))
                        run.font.size = Pt(10)
                        run.font.italic = True
                        run.font.color.rgb = DARK
                        quote_lines = []
                    i += 1
                if quote_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.5)
                    run = p.add_run("\n".join(quote_lines))
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.rgb = DARK
                continue

            if line.startswith("| ") and "---" not in line:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    if "---" not in lines[i]:
                        table_lines.append(lines[i])
                    i += 1
                if table_lines:
                    headers = [c.strip().replace("**", "") for c in table_lines[0].split("|")[1:-1]]
                    rows = []
                    for tl in table_lines[1:]:
                        cols = [c.strip().replace("**", "") for c in tl.split("|")[1:-1]]
                        rows.append(cols)
                    add_table(doc, headers, rows)
                continue

            if line.startswith("- "):
                clean = line[2:].replace("**", "").strip()
                doc.add_paragraph(clean, style="List Bullet")
                i += 1
                continue

            if re.match(r'^\d+\.\s', line):
                clean = re.sub(r'^\d+\.\s*', '', line).replace("**", "").strip()
                doc.add_paragraph(clean, style="List Number")
                i += 1
                continue

            if line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run("\n".join(code_lines))
                    run.font.size = Pt(9)
                    run.font.name = "Consolas"
                    run.font.color.rgb = GRAY
                continue

            if line.startswith("---"):
                i += 1
                continue

            if line == "":
                i += 1
                continue

            clean = line.replace("**", "").strip()
            if clean:
                is_bold = "**" in line and line.count("**") >= 2
                add_body(doc, clean, bold=is_bold)
            i += 1

    docx_path = BASE / "paper.docx"
    doc.save(str(docx_path))
    print(f"Saved DOCX: {docx_path}")

    try:
        from docx2pdf import convert
        pdf_path = BASE / "paper.pdf"
        convert(str(docx_path), str(pdf_path))
        print(f"Saved PDF: {pdf_path}")
    except ImportError:
        print("docx2pdf not available. DOCX created — open in Word and Save As PDF.")
        print(f"Or install: python -m pip install docx2pdf")


if __name__ == "__main__":
    main()
